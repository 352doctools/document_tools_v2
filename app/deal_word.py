# _*_ coding:utf-8 _*_
# Filename: deal_word.py
# Author: pang song
# python 3.6
# Date: 2018/05/08

from docx import  Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches

def typesetting(file_path):
    print("deal docx starting")
    try:
        document = Document(file_path)

        paragraph = document.add_paragraph(u'测试文本')
        # 设置字号
        run = paragraph.add_run(u'24号字')
        run.font.size = Pt(24)
        # 保存文件
        document.save(file_path)
        print("deal docx done!")
        return True
    except Exception as e:
        print(repr(e))
        return False



if __name__ == "__main__":

    #打开文档
    file_path = "C:/PSuse/python_work/stu_python/test.docx"
    document = Document(file_path)

    paragraph = document.add_paragraph(u'添加了文本')

    #加入不同等级的标题
    # document.add_heading(u'标题',0)
    # document.add_heading(u'二级标题',1)
    # document.add_heading(u'二级标题',2)

    #添加文本
    #设置字号
    run = paragraph.add_run(u'设置字号')
    run.font.size=Pt(24)

    #设置字体
    run = paragraph.add_run('Set Font,')
    run.font.name='Consolas'

    #设置中文字体
    run = paragraph.add_run(u'设置中文字体，')
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #设置斜体
    run = paragraph.add_run(u'斜体、')
    run.italic = True

    #设置粗体
    run = paragraph.add_run(u'粗体').bold = True


    #增加表格
    table = document.add_table(rows=3,cols=3)
    hdr_cells=table.rows[0].cells
    hdr_cells[0].text=u"第一列"
    hdr_cells[1].text=u"第二列"
    hdr_cells[2].text=u"第三列"

    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '2'
    hdr_cells[1].text = 'aerszvfdgx'
    hdr_cells[2].text = 'abdzfgxfdf'

    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = '3'
    hdr_cells[1].text = 'cafdwvaef'
    hdr_cells[2].text = 'aabs zfgf'

    #增加分页
    document.add_page_break()

    #保存文件
    document.save('demo.docx')
